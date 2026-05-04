from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"D:\invoice\LedgerAI_Industry_Categories_Spec.docx"


COLORS = {
    "navy": "08294D",
    "teal": "0D9488",
    "blue": "2563EB",
    "amber": "F59E0B",
    "violet": "7C3AED",
    "red": "DC2626",
    "green": "16A34A",
    "ink": "0F172A",
    "muted": "64748B",
    "line": "D8E0EA",
    "bg": "F8FAFC",
    "card": "FFFFFF",
    "soft_teal": "CCFBF1",
    "soft_blue": "DBEAFE",
    "soft_amber": "FEF3C7",
    "soft_violet": "EDE9FE",
}


categories = [
    ("Textile", "Fabric rolls, meters, dye lots, wholesale billing", "Stock by roll/meter, lot tracking, bulk invoices"),
    ("Garment", "Finished clothing, sizes, colors, seasonal collections", "SKU variants, size/color matrix, order fulfillment"),
    ("Electronics", "Devices, accessories, warranty, serial numbers", "Serial tracking, warranty dates, high-value inventory"),
    ("Agriculture", "Seeds, fertilizers, produce, seasonal purchases", "Batch stock, seasonal reports, supplier purchases"),
    ("Mobile Shop", "Phones, accessories, repairs, IMEI tracking", "IMEI/serial fields, service invoices, fast billing"),
    ("E-commerce", "Online orders, returns, multi-channel sales", "Order import-ready layout, returns, shipping status"),
    ("Stationary", "Office supplies, school products, low-value SKUs", "Fast item search, bulk stock updates, reorder alerts"),
    ("Auto Parts", "Vehicle parts, compatibility, mechanics/customers", "Part numbers, vehicle fitment, counter-sale workflow"),
    ("Computer & CCTV", "Hardware, networking, installation services", "Product + service invoices, warranty, AMC reminders"),
    ("Chemical", "Batches, containers, compliance-sensitive stock", "Batch/expiry fields, unit conversion, safety notes"),
    ("IT Company", "Services, subscriptions, projects, retainers", "Service invoices, recurring billing, project reports"),
    ("Marketing Agency", "Campaign services, retainers, client expenses", "Project billing, milestone invoices, expense mapping"),
    ("Packaging", "Boxes, labels, custom sizes, bulk orders", "Custom item dimensions, quotation-to-invoice flow"),
    ("Ceramic", "Tiles, boxes, area coverage, breakage", "Unit conversion, box/sq ft tracking, damage adjustment"),
    ("Furniture", "Made-to-order products, delivery, installation", "Custom orders, advance payments, delivery tracking"),
    ("Building Material", "Cement, steel, sand, transport, large quantities", "Weight/quantity units, transport charges, credit ledger"),
    ("Hardware & Plywood", "Tools, sheets, fittings, mixed retail/wholesale", "Variant stock, cutting/size notes, quick counter billing"),
    ("Cosmetic", "Batch, expiry, MRP, retail stock", "Expiry alerts, MRP vs sale price, fast item lookup"),
    ("Fire Safety", "Equipment, refills, inspections, compliance dates", "Service reminders, certificate tracking, AMC billing"),
    ("Diamond & Jewelry", "High-value items, purity, weight, making charges", "Weight/purity fields, secure access, detailed invoices"),
    ("Scrap", "Weight-based buying/selling, daily rate changes", "Weight slips, rate history, supplier/customer ledger"),
    ("Printing & Designing", "Design services, print quantities, job stages", "Job cards, quotation approvals, production status"),
    ("Rental Services", "Assets, booking period, deposits, returns", "Rental duration billing, deposit ledger, availability"),
    ("Travel Agency", "Bookings, commissions, customers, vendors", "Service vouchers, vendor payable, commission reports"),
]

modules = [
    ("Landing Page", "Show the industry grid and communicate that one app adapts to many business workflows."),
    ("Onboarding", "Ask for business category and preload relevant sample data, fields, reports, and AI prompts."),
    ("Dashboard", "Show common KPIs plus industry-specific risk and opportunity cards."),
    ("Invoices", "Support product, service, mixed, GST-ready, category-field, and AI-drafted invoices."),
    ("Customers", "Maintain customer ledger, balances, notes, contacts, and follow-up history."),
    ("Inventory", "Track SKUs, units, stock value, reorder levels, category fields, and movement."),
    ("Reports", "Provide P&L, sales, purchases, receivables, payables, stock, GST, and category reports."),
    ("User Management", "Control roles, permissions, exports, user invites, and sensitive action audit logs."),
]

ai_examples = [
    ("Textile", "Create invoice for 20 cotton rolls and 12 denim bundles.", "Invoice draft with roll/bundle units and GST."),
    ("Garment", "Which sizes are low before weekend sale?", "Size/color stock risk summary."),
    ("Electronics", "Show warranty items expiring this month.", "Warranty follow-up list."),
    ("Agriculture", "What should I reorder before harvest season?", "Seasonal stock suggestion."),
    ("Mobile Shop", "Create invoice for phone with IMEI and cover.", "Invoice draft with IMEI field."),
    ("E-commerce", "Summarize pending orders and returns.", "Order and return status summary."),
    ("Auto Parts", "Find fast-moving parts for bike customers.", "Stock movement insight."),
    ("Chemical", "Which batches are near expiry?", "Batch expiry alert."),
    ("IT Company", "Create monthly retainer invoice.", "Service invoice draft."),
    ("Marketing Agency", "Summarize client campaign profitability.", "Project margin summary."),
    ("Furniture", "Create advance invoice for custom sofa.", "Advance payment invoice."),
    ("Diamond & Jewelry", "Create invoice with gold weight and making charges.", "High-value detailed invoice draft."),
    ("Rental Services", "Show overdue rental returns.", "Return follow-up list."),
    ("Travel Agency", "Create booking invoice with commission.", "Travel service invoice draft."),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8E0EA"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=None, bold=False, color=None):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, title, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(title)
    style_run(run, 18 if level == 1 else 14, True, COLORS["navy"] if level == 1 else COLORS["ink"])
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(content)
    style_run(run, 10.5, False, COLORS["ink"])
    return p


def add_bullet(doc, content):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(content)
    style_run(run, 10.2, False, COLORS["ink"])
    return p


def style_table(table, widths=None, header_fill="08294D"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_i, row in enumerate(table.rows):
        if r_i == 0:
            set_repeat_table_header(row)
        for c_i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            if widths:
                cell.width = Inches(widths[c_i])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    style_run(run, 9.1 if r_i else 9.4, r_i == 0, COLORS["card"] if r_i == 0 else COLORS["ink"])
            if r_i == 0:
                set_cell_shading(cell, header_fill)
            elif r_i % 2 == 0:
                set_cell_shading(cell, "F8FAFC")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.65)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("LedgerAI Industry Categories Specification")
style_run(r, 26, True, COLORS["navy"])

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Detailed product and design brief for multi-industry invoicing, inventory, customers, reports, user management, and AI Copilot workflows.")
style_run(r, 12, False, COLORS["muted"])

summary = doc.add_table(rows=1, cols=3)
summary.rows[0].cells[0].text = "24 Categories"
summary.rows[0].cells[1].text = "8 Shared Modules"
summary.rows[0].cells[2].text = "AI Copilot Ready"
style_table(summary, [1.9, 1.9, 1.9], "0D9488")

add_heading(doc, "Purpose", 1)
add_body(doc, "LedgerAI should support a wide range of Indian MSME business types while keeping the core product simple: invoicing, customers, inventory, reports, user management, and AI Copilot. Industry selection should influence onboarding, sample data, invoice fields, reports, and AI prompt suggestions.")

add_heading(doc, "Industry Category Matrix", 1)
add_body(doc, "Each category below should be represented in the landing page, onboarding, and demo data. The product should adapt with fields, templates, and AI examples without becoming separate apps.")
add_table(doc, ["Category", "Typical Business Need", "Product Focus"], categories, [1.45, 2.35, 2.55])

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "Shared Product Modules", 1)
add_body(doc, "These modules are common across all industries. Category selection changes labels, fields, examples, and recommended actions, but the core navigation stays consistent.")
add_table(doc, ["Module", "Requirement"], modules, [1.7, 4.8])

add_heading(doc, "Core Inventory Fields", 1)
for item in [
    "Item name, SKU, category, unit, stock quantity, purchase price, sale price, reorder level, and tax rate.",
    "Optional fields: batch number, expiry date, serial number, IMEI, size, color, weight, purity, dimensions, and warranty period.",
    "Inventory should support fast search, low-stock alerts, valuation, purchase suggestions, and stock movement history.",
]:
    add_bullet(doc, item)

add_heading(doc, "Reports Coverage", 1)
for item in [
    "Common reports: profit and loss, sales summary, purchase summary, receivables, payables, inventory detail, stock movement, GST summary, and expense report.",
    "Category reports: expiry report, serial/warranty report, size/color stock report, weight-based stock report, rental utilization report, and project profitability report.",
    "AI should explain reports in plain language and suggest the next business action.",
]:
    add_bullet(doc, item)

add_heading(doc, "User Management Requirements", 1)
add_table(doc, ["Role", "Access Pattern"], [
    ("Owner", "All access, billing, settings, users, exports, and audit logs."),
    ("Admin", "Daily operations, users, invoices, stock, customers, and reports."),
    ("Accountant", "Invoices, payments, reports, taxes, customers, and exports."),
    ("Sales", "Customers, quotations, invoices, and payment reminders."),
    ("Inventory Manager", "Items, stock movement, reorder alerts, and purchase needs."),
    ("Viewer", "Read-only access to assigned areas."),
], [1.7, 4.8])

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "AI Copilot Requirements", 1)
add_body(doc, "AI Copilot should work across all categories with deterministic demo responses first. It should help users draft invoices, summarize receivables, identify risks, suggest reorder quantities, explain reports, and draft customer follow-ups.")

add_heading(doc, "Shared AI Actions", 2)
for item in [
    "Draft invoice from plain English.",
    "Summarize sales, receivables, expenses, and stock.",
    "Identify overdue customers and low-stock items.",
    "Suggest reorder quantities and payment follow-ups.",
    "Explain reports in plain language.",
    "Draft WhatsApp or email reminder text for user review.",
]:
    add_bullet(doc, item)

add_heading(doc, "Category-Specific AI Examples", 2)
add_table(doc, ["Category", "AI Prompt Example", "Expected AI Output"], ai_examples, [1.35, 2.65, 2.55])

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "Design Requirements", 1)
for item in [
    "Use Tailwind-friendly design tokens with semantic colors.",
    "Default theme follows system preference; users can override with light or dark mode.",
    "Use a clean SaaS dashboard style with dense but readable operational tables.",
    "Industry category grid should use consistent icons, spacing, and labels.",
    "Use editable 3D vector visuals for landing hero, empty states, and AI onboarding.",
    "Avoid copied Hisab branding, icons, images, or long-form text.",
]:
    add_bullet(doc, item)

add_heading(doc, "Full Workflow Coverage", 1)
workflow = [
    "Landing page with industry category grid.",
    "Business category selection during onboarding.",
    "Dashboard customized by selected category.",
    "Invoice list, AI invoice builder, invoice detail, and preview.",
    "Customer list plus customer ledger/profile.",
    "Inventory list plus item detail with category-specific fields.",
    "Reports dashboard plus report detail with AI summary.",
    "User management list plus role and permissions editor.",
    "Company settings and theme settings: system, light, dark.",
    "Mobile dashboard and mobile invoice creation flow.",
]
for item in workflow:
    add_bullet(doc, item)

add_heading(doc, "Acceptance Criteria", 1)
for item in [
    "All 24 listed industry categories are represented.",
    "Each category has a clear business reason and product behavior.",
    "Core modules support category-specific fields without creating separate apps.",
    "AI Copilot has shared and category-specific prompt examples.",
    "The design supports desktop and mobile workflows.",
    "The design supports system-default, light, and dark theme modes.",
    "This document can be used directly as a product/design brief for Figma and React implementation.",
]:
    add_bullet(doc, item)

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("LedgerAI Industry Categories Specification")
    style_run(run, 8.5, False, COLORS["muted"])

doc.save(OUT)
print(OUT)
